"""
Generates INVOICE_PROCESSING_PROMPT.docx from INVOICE_PROCESSING_PROMPT.md.
Run after any edit to the .md file:
    python3 build_prompt_docx.py

Structure of the output:
  • Pages 1-N  : General rules (intro, JSON structure, calculation rules, checklist)
  • Then       : One dedicated page per supplier (coloured banner + structured sections)
"""

import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = os.path.join(os.path.dirname(__file__), "INVOICE_PROCESSING_PROMPT.md")
DEST = os.path.join(os.path.dirname(__file__), "INVOICE_PROCESSING_PROMPT.docx")

# ── Supplier banner colours (one per supplier, in order of appearance) ─────────
BANNER_COLORS = [
    ("1a56db", "ffffff"),  # QBL            – blue
    ("b91c1c", "ffffff"),  # VKS             – red
    ("0369a1", "ffffff"),  # Edendale        – sky
    ("6d28d9", "ffffff"),  # Innodis         – violet
    ("c2410c", "ffffff"),  # RR Rapid        – orange
    ("065f46", "ffffff"),  # BrandActiv/IBL  – forest green
    ("3730a3", "ffffff"),  # Lim How         – indigo
    ("9d174d", "ffffff"),  # Cadi Fortune    – rose
    ("1e40af", "ffffff"),  # Panagora        – deep blue
    ("0f766e", "ffffff"),  # Kool Food       – teal
    ("075985", "ffffff"),  # Pillay R        – dark sky
    ("166534", "ffffff"),  # Li Wan Po       – dark green
    ("4c1d95", "ffffff"),  # Intl Distillers – purple
    ("78350f", "ffffff"),  # Tea Blenders    – brown
    ("1f2937", "ffffff"),  # Freelance       – charcoal
]

# ── XML helpers ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6)
    tcPr.append(shd)

def set_table_borders(table, color="DDDDDD", sz="4"):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)

def no_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)

def set_para_bg(para, hex6):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6)
    pPr.append(shd)

def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)

# ── Run helpers ────────────────────────────────────────────────────────────────
def add_run(para, text, bold=False, italic=False, size=None, color=None, font=None, code=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if code:
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0xc7, 0x25, 0x3e)
    if font:
        run.font.name = font
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color.lstrip("#")))
    return run

def parse_inline(para, text, base_size=None):
    """Render **bold**, *italic*, `code` inline markers."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    for part in pattern.split(text):
        if not part: continue
        if part.startswith("**") and part.endswith("**"):
            r = add_run(para, part[2:-2], bold=True)
        elif part.startswith("*") and part.endswith("*"):
            r = add_run(para, part[1:-1], italic=True)
        elif part.startswith("`") and part.endswith("`"):
            r = add_run(para, part[1:-1], code=True)
        else:
            r = add_run(para, part)
        if base_size:
            r.font.size = Pt(base_size)

# ── General-rules renderer ─────────────────────────────────────────────────────
def render_general(doc, lines):
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Skip top meta comment lines
        if line.startswith("# ─") or (line.startswith("# ") and ("Copy everything" in line or "Paste everything" in line)):
            i += 1; continue
        if line.strip() == "---": i += 1; continue

        # H1
        if line.startswith("# ") and not line.startswith("# ─"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(line[2:])
            r.bold = True; r.font.size = Pt(18)
            r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
            i += 1; continue

        # H2
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(line[3:])
            r.bold = True; r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0x0e, 0x7c, 0x86)
            i += 1; continue

        # H3
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(line[4:])
            r.bold = True; r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x10, 0x6b, 0x3d)
            i += 1; continue

        # Code block
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run("\n".join(code_lines))
            r.font.name  = "Courier New"
            r.font.size  = Pt(9)
            r.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
            set_para_bg(p, "F3F4F6")
            continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            rows = [r2 for r2 in table_lines if not re.match(r"^\s*\|[\s\-|:]+\|\s*$", r2)]
            if not rows: continue
            parsed = [[c.strip() for c in r2.strip().strip("|").split("|")] for r2 in rows]
            num_cols = max(len(r2) for r2 in parsed)
            tbl = doc.add_table(rows=len(parsed), cols=num_cols)
            set_table_borders(tbl)
            for ri, row_data in enumerate(parsed):
                for ci, ct in enumerate(row_data):
                    if ci >= num_cols: break
                    cell = tbl.cell(ri, ci)
                    cell.paragraphs[0].clear()
                    p2 = cell.paragraphs[0]
                    p2.paragraph_format.space_before = Pt(2)
                    p2.paragraph_format.space_after  = Pt(2)
                    parse_inline(p2, ct, base_size=10)
                    if ri == 0:
                        set_cell_bg(cell, "DCE6F1")
                        for run in p2.runs: run.bold = True
            doc.add_paragraph()
            continue

        # Checkbox
        if line.startswith("- [ ]"):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            parse_inline(p, "☐  " + line[6:].strip())
            i += 1; continue

        # Bullet
        if re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            parse_inline(p, line[2:].strip())
            i += 1; continue

        # Numbered list
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            parse_inline(p, m.group(1))
            i += 1; continue

        # Blank
        if line.strip() == "":
            i += 1; continue

        # Normal
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        parse_inline(p, line)
        i += 1

# ── Supplier page renderer ─────────────────────────────────────────────────────
def section_label(doc, text, bg_hex, text_hex="ffffff"):
    """Full-width coloured label strip for a section heading."""
    tbl = doc.add_table(rows=1, cols=1)
    no_borders(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(*bytes.fromhex(text_hex))
    return tbl

def render_supplier_page(doc, supplier_text, banner_color, text_color, supplier_idx, total):
    """Render one full supplier page."""
    doc.add_page_break()

    # ── Parse supplier block into named sections ───────────────────────────────
    lines = supplier_text.split("\n")

    # First line is the heading "### Supplier: NAME" — already stripped
    name_line = lines[0].strip() if lines else "Unknown"
    name = re.sub(r"^#+\s*Supplier:\s*", "", name_line).strip()
    body_lines = lines[1:]

    fields = {
        "category": "",
        "supabase_name": "",
        "columns": [],
        "colorder": "",
        "vat": [],
        "totals_keys": "",
        "special": [],
        "gotchas": [],
        "extra_notes": [],
    }

    current = None
    col_table_lines = []
    in_col_table = False

    for line in body_lines:
        l = line.strip()
        if l.startswith("**Category:**"):
            fields["category"] = l.replace("**Category:**", "").strip()
            current = None
        elif l.startswith("**Supabase supplier_name:**") or l.startswith("**Supabase meta:**"):
            fields["supabase_name"] += l + " "
            current = None
        elif l.startswith("**Invoice columns"):
            current = "columns"
            in_col_table = False
        elif l.startswith("**colOrder:**"):
            fields["colorder"] = re.sub(r"\*\*colOrder:\*\*\s*", "", l).strip().strip("`")
            current = None
        elif l.startswith("**VAT treatment:**") or l.startswith("**VAT Treatment:**"):
            current = "vat"
        elif l.startswith("**Totals keys"):
            fields["totals_keys"] = l
            current = None
        elif l.startswith("**Special rules:**") or l.startswith("**Special Rules:**"):
            current = "special"
        elif l.startswith("**Gotchas:**"):
            current = "gotchas"
        elif l == "" or l == "---":
            pass
        else:
            if current == "columns":
                if "|" in l and l.startswith("|"):
                    col_table_lines.append(l)
            elif current == "vat":
                if re.match(r"^[-*]|\d+\.", l):
                    fields["vat"].append(re.sub(r"^[-*\d.]+\s*", "", l))
            elif current == "special":
                if re.match(r"^[-*]|\d+\.", l):
                    fields["special"].append(re.sub(r"^[-*\d.]+\s*", "", l))
            elif current == "gotchas":
                if re.match(r"^[-*]|\d+\.", l):
                    fields["gotchas"].append(re.sub(r"^[-*\d.]+\s*", "", l))

    # Parse column table
    col_rows = []
    if col_table_lines:
        clean = [r for r in col_table_lines if not re.match(r"^\s*\|[\s\-|:]+\|\s*$", r)]
        for r in clean:
            cells = [c.strip() for c in r.strip().strip("|").split("|")]
            col_rows.append(cells)

    # ── Banner ──────────────────────────────────────────────────────────────────
    banner_tbl = doc.add_table(rows=1, cols=1)
    no_borders(banner_tbl)
    banner_cell = banner_tbl.cell(0, 0)
    set_cell_bg(banner_cell, banner_color)
    banner_p = banner_cell.paragraphs[0]
    banner_p.paragraph_format.space_before = Pt(8)
    banner_p.paragraph_format.space_after  = Pt(4)
    banner_p.paragraph_format.left_indent  = Cm(0.4)
    r = banner_p.add_run(name.upper())
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(*bytes.fromhex(text_color))

    if fields["category"]:
        cat_p = banner_cell.add_paragraph()
        cat_p.paragraph_format.space_before = Pt(0)
        cat_p.paragraph_format.space_after  = Pt(8)
        cat_p.paragraph_format.left_indent  = Cm(0.4)
        cr = cat_p.add_run(fields["category"])
        cr.italic = True
        cr.font.size = Pt(11)
        cr.font.color.rgb = RGBColor(*bytes.fromhex(text_color))
        cr.font.color.theme_color = None

    if fields["supabase_name"].strip():
        sn_p = banner_cell.add_paragraph()
        sn_p.paragraph_format.space_before = Pt(0)
        sn_p.paragraph_format.space_after  = Pt(8)
        sn_p.paragraph_format.left_indent  = Cm(0.4)
        sr2 = sn_p.add_run(fields["supabase_name"].strip())
        sr2.font.size = Pt(9)
        sr2.font.color.rgb = RGBColor(*bytes.fromhex(text_color))

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Columns table ───────────────────────────────────────────────────────────
    if col_rows:
        section_label(doc, "  INVOICE COLUMNS  (left → right on paper)", banner_color)
        num_cols = max(len(r) for r in col_rows)
        tbl = doc.add_table(rows=len(col_rows), cols=num_cols)
        set_table_borders(tbl, color="CCCCCC")
        for ri, row_data in enumerate(col_rows):
            for ci, ct in enumerate(row_data):
                if ci >= num_cols: break
                cell = tbl.cell(ri, ci)
                cell.paragraphs[0].clear()
                p2 = cell.paragraphs[0]
                p2.paragraph_format.space_before = Pt(3)
                p2.paragraph_format.space_after  = Pt(3)
                p2.paragraph_format.left_indent  = Cm(0.15)
                parse_inline(p2, ct, base_size=10)
                if ri == 0:
                    set_cell_bg(cell, "E8EDF5")
                    for run in p2.runs: run.bold = True
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(6)

    # ── colOrder ────────────────────────────────────────────────────────────────
    if fields["colorder"]:
        section_label(doc, "  colOrder  —  paste this into the invoice JSON", banner_color)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.2)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(fields["colorder"])
        r.font.name  = "Courier New"
        r.font.size  = Pt(10)
        r.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
        set_para_bg(p, "F3F4F6")

    # ── VAT treatment ───────────────────────────────────────────────────────────
    if fields["vat"]:
        section_label(doc, "  VAT TREATMENT", "2d6a4f")
        for item in fields["vat"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            parse_inline(p, "• " + item, base_size=10)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if fields["totals_keys"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_after  = Pt(4)
        parse_inline(p, fields["totals_keys"], base_size=10)

    # ── Special rules ───────────────────────────────────────────────────────────
    if fields["special"]:
        section_label(doc, "  SPECIAL RULES", "1d4ed8")
        for idx, item in enumerate(fields["special"], 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            parse_inline(p, f"{idx}.  " + item, base_size=10)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Gotchas ─────────────────────────────────────────────────────────────────
    if fields["gotchas"]:
        section_label(doc, "  ⚠  GOTCHAS", "b45309")
        for item in fields["gotchas"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            parse_inline(p, "⚠  " + item, base_size=10)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Supplier counter at bottom
    counter_p = doc.add_paragraph()
    counter_p.paragraph_format.space_before = Pt(8)
    cr = counter_p.add_run(f"Supplier {supplier_idx} of {total}")
    cr.font.size = Pt(8)
    cr.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
    cr.italic = True


# ── Main ───────────────────────────────────────────────────────────────────────
def build():
    with open(SRC, encoding="utf-8") as f:
        content = f.read()

    # Split into general rules + supplier sections
    supplier_marker = "\n## Supplier-Specific Rules"
    if supplier_marker in content:
        general_text, supplier_text = content.split(supplier_marker, 1)
    else:
        general_text = content
        supplier_text = ""

    # Parse individual supplier blocks
    supplier_blocks = []
    if supplier_text:
        parts = re.split(r"\n(?=### Supplier:)", supplier_text)
        for part in parts:
            part = part.strip()
            if part.startswith("### Supplier:"):
                supplier_blocks.append(part)

    # Build document
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # General rules
    general_lines = general_text.splitlines(keepends=True)
    render_general(doc, general_lines)

    # Supplier pages
    total = len(supplier_blocks)
    for idx, block in enumerate(supplier_blocks):
        color_pair = BANNER_COLORS[idx % len(BANNER_COLORS)]
        render_supplier_page(doc, block, color_pair[0], color_pair[1], idx + 1, total)

    doc.save(DEST)
    print(f"✓  Saved {total} supplier pages → {DEST}")

if __name__ == "__main__":
    build()
